import base64
import os
import requests
import csv
from bs4 import BeautifulSoup, Tag, NavigableString
from docx import Document
from docx.shared import Inches, Pt

# API endpoint and token setup
API_URL = 'https://alleducationboardresults.com/api/blogs/'  # Replace with your actual API endpoint
API_TOKEN = 'f4017f63b5182786d456cadf8844b6a4438f9342'    # Replace with your actual API token

HEADERS = {
    'Authorization': f'Token {API_TOKEN}'
}

# Directory to temporarily save images
TEMP_IMAGE_DIR = 'temp_images'
# Directory to save documents
DOCUMENTS_DIR = 'documents'

# Ensure the directories exist
if not os.path.exists(TEMP_IMAGE_DIR):
    os.makedirs(TEMP_IMAGE_DIR)
if not os.path.exists(DOCUMENTS_DIR):
    os.makedirs(DOCUMENTS_DIR)




# ----------------------- Helper Functions -----------------------

def add_spacing(doc, inches=0.2):
    """
    Add a blank paragraph with specified spacing after an element.

    Args:
        doc (Document): The DOCX document object.
        inches (float): The amount of space to add in inches. Default is 0.2 inches.
    """
    try:
        paragraph = doc.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(inches * 72)  # 1 inch = 72 points
        paragraph_format.space_after = Pt(0)
    except Exception:
        pass  # Silent error handling

def handle_base64_image(data_url):
    """
    Decode a base64 image from a data URL and save it to the temporary directory.

    Args:
        data_url (str): The data URL containing the base64-encoded image.

    Returns:
        str: The file path to the saved image, or None if failed.
    """
    try:
        header, encoded = data_url.split(",", 1)
        data = base64.b64decode(encoded)
        file_extension = header.split('/')[1].split(';')[0]  # e.g., 'png' from 'data:image/png;base64,...'
        filename = f'image_{len(os.listdir(TEMP_IMAGE_DIR)) + 1}.{file_extension}'
        file_path = os.path.join(TEMP_IMAGE_DIR, filename)
        with open(file_path, 'wb') as file:
            file.write(data)
        return file_path
    except Exception:
        return None

def download_image(url):
    """
    Download an image from a URL and save it to the temporary directory.

    Args:
        url (str): The URL of the image to download.

    Returns:
        str: The file path to the saved image, or None if failed.
    """
    try:
        response = requests.get(url, stream=True)
        if response.status_code == 200:
            filename = os.path.basename(url.split('?')[0])  # Remove query parameters
            if not filename:
                filename = f'image_{len(os.listdir(TEMP_IMAGE_DIR)) + 1}.png'  # Default to png if no name
            file_path = os.path.join(TEMP_IMAGE_DIR, filename)
            with open(file_path, 'wb') as f:
                for chunk in response.iter_content(1024):
                    f.write(chunk)
            return file_path
        else:
            return None
    except Exception:
        return None

def add_table(doc, soup_table):
    """
    Convert an HTML table to a DOCX table and add it to the document.

    Args:
        doc (Document): The DOCX document object.
        soup_table (Tag): The BeautifulSoup Tag object representing the table.
    """
    try:
        rows = soup_table.find_all('tr')
        if not rows:
            return

        # Determine number of columns based on the first row
        first_row = rows[0]
        cols = first_row.find_all(['th', 'td'])
        num_cols = len(cols)

        # Create a DOCX table with the number of columns
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'

        # Add header cells if present
        header_cells = first_row.find_all('th')
        if header_cells:
            for idx, cell in enumerate(header_cells):
                table.cell(0, idx).text = cell.get_text(strip=True)
            data_rows = rows[1:]
        else:
            # If no header, treat all rows as data
            data_rows = rows

        # Add data rows
        for row in data_rows:
            cells = row.find_all(['th', 'td'])
            if len(cells) == 0:
                continue  # Skip empty rows
            row_cells = table.add_row().cells
            for idx, cell in enumerate(cells):
                if idx < num_cols:
                    row_cells[idx].text = cell.get_text(strip=True)
        # Add spacing after the table
        add_spacing(doc)
    except Exception:
        pass  # Silent error handling

def extract_list_items(list_element):
    """
    Recursively extract list items from a <ul> or <ol> element.

    Args:
        list_element (Tag): BeautifulSoup Tag object representing <ul> or <ol>.

    Returns:
        list: A nested list representing the list structure.
    """
    items = []
    for li in list_element.find_all('li', recursive=False):
        # Initialize a list to hold the item's content
        item_content = []

        # Iterate over the children of the <li>
        for child in li.children:
            if isinstance(child, NavigableString):
                text = child.strip()
                if text:
                    item_content.append(text)
            elif isinstance(child, Tag):
                if child.name in ['ul', 'ol']:
                    # Recursively extract items from the nested list
                    nested_items = extract_list_items(child)
                    item_content.append(nested_items)
                else:
                    # For other tags like <p>, <img>, etc.
                    text = child.get_text(strip=True)
                    if text:
                        item_content.append(text)

        # Combine the main text and any nested lists
        main_text = ''
        nested_lists = []
        for content in item_content:
            if isinstance(content, list):
                nested_lists.append(content)
            else:
                main_text += content + ' '

        main_text = main_text.strip()

        if nested_lists:
            # If there are nested lists, append them as sublists
            for nested in nested_lists:
                items.append([main_text, nested])
        else:
            # Simple list item
            items.append(main_text)
    return items

def add_list(doc, items, list_type='ul', level=0):
    """
    Add a list (unordered or ordered) to the DOCX document.

    Args:
        doc (Document): The DOCX document object.
        items (list): A list of strings or nested lists representing list items.
        list_type (str): Type of list - 'ul' for unordered, 'ol' for ordered.
        level (int): Indentation level for nested lists.
    """
    try:
        for item in items:
            if isinstance(item, list):
                # The first element is the main text, the second is the nested list
                main_text, nested_items = item
                # Add the main list item
                if list_type == 'ul':
                    style = 'List Bullet'
                elif list_type == 'ol':
                    style = 'List Number'
                else:
                    style = 'Normal'  # Fallback to normal style

                paragraph = doc.add_paragraph(main_text, style=style)
                if level > 0:
                    # Indent the list item based on its level
                    paragraph.paragraph_format.left_indent = Inches(0.25 * level)

                # Recursively add the nested list
                add_list(doc, nested_items, list_type=list_type, level=level+1)
            else:
                # Simple list item
                if list_type == 'ul':
                    style = 'List Bullet'
                elif list_type == 'ol':
                    style = 'List Number'
                else:
                    style = 'Normal'  # Fallback to normal style

                paragraph = doc.add_paragraph(item, style=style)
                if level > 0:
                    # Indent the list item based on its level
                    paragraph.paragraph_format.left_indent = Inches(0.25 * level)
        # Add spacing after the list
        add_spacing(doc)
    except Exception:
        pass  # Silent error handling

def add_heading(doc, text, level):
    """
    Add a heading to the DOCX document with appropriate formatting.

    Args:
        doc (Document): The DOCX document object.
        text (str): The heading text.
        level (int): The heading level (1 for h1, 2 for h2, ..., 6 for h6).
    """
    try:
        # Map HTML heading level directly to python-docx heading level
        heading_level = max(0, min(level, 8))  # python-docx supports up to Heading 9

        # Add the heading with the correct level
        heading = doc.add_heading(text, level=heading_level)

        # Access the first run in the heading paragraph to apply formatting
        if heading.runs:
            run = heading.runs[0]
            run.bold = True  # Ensure the heading is bold

        # Add spacing after the heading
        add_spacing(doc)
    except Exception:
        pass  # Silent error handling

def add_paragraph(doc, text):
    """
    Add a paragraph to the DOCX document.

    Args:
        doc (Document): The DOCX document object.
        text (str): The paragraph text.
    """
    try:
        doc.add_paragraph(text)
    except Exception:
        pass  # Silent error handling

def add_image(doc, src):
    """
    Add an image to the DOCX document, handling both base64 and URL images.

    Args:
        doc (Document): The DOCX document object.
        src (str): The source of the image (data URL or image URL).
    """
    try:
        if src.startswith('data:'):
            image_path = handle_base64_image(src)
        elif src:
            image_path = download_image(src)
        else:
            return

        if image_path and os.path.exists(image_path):
            doc.add_picture(image_path, width=Inches(5))
            os.remove(image_path)  # Remove the image after adding to the document
            # Add spacing after the image
            add_spacing(doc)
        else:
            return
    except Exception:
        pass  # Silent error handling

def process_element(doc, element):
    """
    Recursively process an HTML element and add corresponding content to the DOCX document.

    Args:
        doc (Document): The DOCX document object.
        element (Tag): The BeautifulSoup Tag object to process.
    """
    if isinstance(element, NavigableString):
        # Ignore standalone strings
        return
    elif not isinstance(element, Tag):
        return

    if element.name in ['p', 'div']:
        text = element.get_text(strip=True)
        if text:
            add_paragraph(doc, text)
    elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(element.name[1:])
        text = element.get_text(strip=True)
        if text:
            add_heading(doc, text, level)
    elif element.name in ['ul', 'ol']:
        list_type = 'ul' if element.name == 'ul' else 'ol'
        items = extract_list_items(element)
        add_list(doc, items, list_type=list_type)
        return  # Prevent processing child elements again
    elif element.name == 'img':
        src = element.get('src', '')
        add_image(doc, src)
    elif element.name == 'table':
        add_table(doc, element)
        return  # Prevent processing child elements again
    # Add more tags here if needed (e.g., blockquotes, code blocks)

    # Recursively process the next level of children
    for child in element.children:
        process_element(doc, child)

def html_to_docx(html_content, filename):
    """
    Convert HTML content to a DOCX document.

    Args:
        html_content (str): The HTML content as a string.
        filename (str): The desired filename for the DOCX document.

    Returns:
        str: Confirmation message.
    """
    try:
        doc = Document()
        soup = BeautifulSoup(html_content, 'html.parser')

        # Use <body> if present, else use the entire soup
        content = soup.body if soup.body else soup

        # Start processing from the immediate children of content
        for element in content.find_all(recursive=False):
            process_element(doc, element)

        doc.save(DOCUMENTS_DIR + '/' + filename)
        return f"Document saved as '{filename}'"
    except Exception:
        return "Conversion failed due to an error."

def get_blog_post_and_convert_to_docx(post_id):
    """
    Fetch a blog post by ID from the API and convert its HTML content to DOCX.

    Args:
        post_id (str): The ID of the blog post.

    Returns:
        str: Confirmation message or error message.
    """
    try:
        response = requests.get(f"{API_URL}{post_id}/", headers=HEADERS)
        if response.status_code == 200:
            data = response.json()
            html_content = data.get('content')
            if html_content:
                filename = f'{post_id}.docx'
                return html_to_docx(html_content, filename)
            else:
                return "No content available for this post."
        else:
            return f"Failed to fetch post: Status code {response.status_code}"
    except Exception:
        return "Failed to fetch post due to an exception."

def cleanup_temp_images():
    """
    Remove all files in the temporary image directory and delete the directory.
    """
    try:
        for filename in os.listdir(TEMP_IMAGE_DIR):
            file_path = os.path.join(TEMP_IMAGE_DIR, filename)
            if os.path.isfile(file_path):
                os.remove(file_path)
        os.rmdir(TEMP_IMAGE_DIR)
    except Exception:
        pass  # Silent error handling



if __name__ == "__main__":
    csv_file = input("Enter the path to your CSV file: ").strip()
    if not os.path.exists(csv_file):
        print(f"[ERROR] File not found: {csv_file}")
    with open(csv_file, 'r', encoding='utf-8') as file:
        reader = csv.DictReader(file)
        for row in reader:
            post_id = row.get('ID')
            if post_id:
                print(f"Processing post ID: {post_id}")
                result = get_blog_post_and_convert_to_docx(post_id)
                print(result)

    print("All posts processed.")